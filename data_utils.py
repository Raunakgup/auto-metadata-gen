# data_utils.py

import os
from datetime import datetime
from typing import Dict

from docx import Document
from docx import Document as DocxDocument
from PyPDF2 import PdfReader
from PyPDF2 import PdfReader as PDFReader
from pdf2image import convert_from_path
import pytesseract


def extract_text_from_txt(path: str) -> str:
    """Read plain-text files."""
    with open(path, 'r', encoding='utf-8', errors='ignore') as f:
        return f.read()


def extract_text_from_docx(path: str) -> str:
    """Extract text from a .docx file."""
    doc = Document(path)
    return "\n".join(para.text for para in doc.paragraphs)


def extract_text_from_pdf(path: str) -> str:
    """
    Extract text from a PDF.
    1) Try PyPDF2 text extraction.
    2) If that yields very little text, run OCR via pytesseract.
    """
    reader = PdfReader(path)
    text_pages = [page.extract_text() or "" for page in reader.pages]
    joined = "\n".join(text_pages).strip()

    # Fallback to OCR if very little text
    if len(joined) < 100:
        ocr_pages = []
        images = convert_from_path(path, dpi=300)
        for img in images:
            ocr_pages.append(pytesseract.image_to_string(img))
        joined = "\n".join(ocr_pages)
    return joined


def extract_text(path: str) -> str:
    """
    Dispatcher: choose extractor based on file extension.
    """
    ext = os.path.splitext(path)[1].lower()
    if ext == ".txt":
        return extract_text_from_txt(path)
    elif ext == ".docx":
        return extract_text_from_docx(path)
    elif ext == ".pdf":
        return extract_text_from_pdf(path)
    else:
        raise ValueError(f"Unsupported file type: {ext}")


def get_pdf_metadata(path: str) -> Dict[str, str]:
    """Extract author and creation date from a PDF, if present."""
    reader = PDFReader(path)
    info = reader.metadata or {}

    # Handle both attribute‐style and dict‐style metadata
    if hasattr(info, "author"):
        author = info.author or ""
    else:
        # PDF metadata dict often uses keys like '/Author'
        author = info.get("author") or info.get("/Author") or ""

    # Creation date can be attribute or dict entry
    raw = None
    if hasattr(info, "creation_date"):
        raw = info.creation_date
    else:
        raw = info.get("creation_date") or info.get("/CreationDate")

    created_at = ""
    if raw:
        if isinstance(raw, datetime):
            created_at = raw.isoformat()
        else:
            s = str(raw)
            if s.startswith("D:"):
                s = s[2:]
            try:
                created_at = datetime.strptime(s[:14], "%Y%m%d%H%M%S").isoformat()
            except Exception:
                created_at = s
    return {"author": author, "created_at": created_at}


def get_docx_metadata(path: str) -> Dict[str, str]:
    """Extract author and creation date from a DOCX, if present."""
    doc = DocxDocument(path)
    props = doc.core_properties
    author = props.author or ""
    created_at = ""
    if props.created:
        # often a datetime object
        if isinstance(props.created, datetime):
            created_at = props.created.isoformat()
        else:
            created_at = str(props.created)
    return {"author": author, "created_at": created_at}


def get_file_metadata(path: str) -> Dict[str, str]:
    """
    Dispatcher for embedded file metadata based on extension.
    Returns dict with 'author' and 'created_at' keys (always strings).
    """
    ext = os.path.splitext(path)[1].lower()
    if ext == ".pdf":
        return get_pdf_metadata(path)
    elif ext == ".docx":
        return get_docx_metadata(path)
    else:
        return {"author": "", "created_at": ""}


if __name__ == "__main__":
    # Quick CLI test
    import sys
    if len(sys.argv) != 2:
        print("Usage: python data_utils.py <path_to_file>")
        sys.exit(1)

    p = sys.argv[1]
    print("== Text Preview ==")
    print(extract_text(p)[:500])
    print("\n== Embedded Metadata ==")
    print(get_file_metadata(p))
